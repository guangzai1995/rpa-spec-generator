"""Word 文档生成器：使用 docxtpl 渲染 RPA 需求规格说明书"""
import base64
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple, Union

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from PIL import Image as PILImage, ImageDraw, ImageFont
from openai import OpenAI
import httpx

from app.models.schemas import ExtractionResult
from app.utils.logger import get_logger

logger = get_logger(__name__)

TEMPLATE_DIR = Path(__file__).parent.parent.parent / "templates"
OUTPUT_DIR = Path(os.getenv("STATIC_DIR", "static")) / "docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _normalize_image(path: str) -> io.BytesIO:
    """用 PIL 重新编码图片，确保 python-docx 能识别格式"""
    img = PILImage.open(path).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _parse_bbox_coords(text: str) -> Optional[List[int]]:
    """从 VLM 响应中解析坐标（支持多种格式）。返回 [x1,y1,x2,y2] 或 None。"""
    clean = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
    # 格式: <box>(x1,y1),(x2,y2)</box>
    m = re.search(r'<box>\((\d+),\s*(\d+)\),?\s*\((\d+),\s*(\d+)\)</box>', clean)
    if m:
        return [int(m.group(i)) for i in range(1, 5)]
    # 格式: <box>x1,y1,x2,y2</box>
    m = re.search(r'<box>(\d+),\s*(\d+),\s*(\d+),\s*(\d+)</box>', clean)
    if m:
        return [int(m.group(i)) for i in range(1, 5)]
    # 格式: [x1,y1,x2,y2]
    m = re.search(r'\[(\d+),\s*(\d+),\s*(\d+),\s*(\d+)\]', clean)
    if m:
        return [int(m.group(i)) for i in range(1, 5)]
    # 格式: (x1,y1,x2,y2)
    m = re.search(r'\((\d+),\s*(\d+),\s*(\d+),\s*(\d+)\)', clean)
    if m:
        return [int(m.group(i)) for i in range(1, 5)]
    return None


def _detect_ui_targets(frame_path: str, targets: List[str]) -> List[dict]:
    """调用视觉 LLM 检测截图中 UI 元素的边界框坐标。

    使用 Qwen VL 原生 <ref> 格式 prompt，坐标为千分比(0-1000)自动转为像素值。
    双策略：先用 ref+box prompt，失败则用 /no_think+ref 重试。

    Args:
        frame_path: 截图文件路径
        targets: 要检测的 UI 目标列表，如 ["导出按钮", "项目编号搜索框"]

    Returns:
        [{"target": "导出按钮", "bbox": [x1, y1, x2, y2]}, ...]
        坐标为像素值
    """
    vision_api_key = os.getenv("VISION_API_KEY", "") or os.getenv("LLM_API_KEY", "")
    vision_base_url = os.getenv("VISION_BASE_URL", "") or os.getenv("LLM_BASE_URL", "")
    vision_model = os.getenv("VISION_MODEL", "") or os.getenv("LLM_MODEL", "")

    if not all([vision_api_key, vision_base_url, vision_model]):
        return []

    # 双策略 prompt（按优先级）
    prompt_strategies = [
        '请用<ref>{target}</ref>找到图中对应元素的位置，输出<box>坐标</box>。',
        '/no_think\n<ref>{target}</ref>',
    ]

    try:
        img = PILImage.open(frame_path).convert("RGB")
        orig_w, orig_h = img.size
        max_size = 1024
        if max(orig_w, orig_h) > max_size:
            ratio = max_size / max(orig_w, orig_h)
            img = img.resize((int(orig_w * ratio), int(orig_h * ratio)), PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        client = OpenAI(
            api_key=vision_api_key,
            base_url=vision_base_url,
            timeout=httpx.Timeout(30.0, connect=10.0),
        )

        valid = []
        for target in targets[:3]:
            coords = None
            for prompt_tpl in prompt_strategies:
                prompt = prompt_tpl.format(target=target)
                response = client.chat.completions.create(
                    model=vision_model,
                    messages=[{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {
                                "url": f"data:image/jpeg;base64,{b64}", "detail": "high"
                            }},
                        ],
                    }],
                    temperature=0.0,
                    max_tokens=256,
                )
                content = (response.choices[0].message.content or "").strip()
                coords = _parse_bbox_coords(content)
                if coords and coords[0] < coords[2] and coords[1] < coords[3]:
                    break  # 成功，跳过备选策略
                coords = None  # 重置，尝试下一策略

            if coords:
                # 千分比坐标 (0-1000) → 像素坐标
                px = [
                    int(coords[0] * orig_w / 1000),
                    int(coords[1] * orig_h / 1000),
                    int(coords[2] * orig_w / 1000),
                    int(coords[3] * orig_h / 1000),
                ]
                # 钳位到图片范围
                px[0] = max(0, min(px[0], orig_w))
                px[1] = max(0, min(px[1], orig_h))
                px[2] = max(0, min(px[2], orig_w))
                px[3] = max(0, min(px[3], orig_h))
                if px[2] - px[0] >= 3 and px[3] - px[1] >= 3:
                    valid.append({"target": target, "bbox": px})
                    logger.debug(f"  bbox [{target}]: ‰{coords} → px{px}")

        if valid:
            logger.info(f"UI目标检测: 找到 {len(valid)}/{len(targets)} 个元素")
        return valid
    except Exception as e:
        logger.warning(f"UI目标检测失败: {e}")
        return []


def _annotate_frame(frame_path: str, proc, frame_descriptions: Optional[dict] = None) -> io.BytesIO:
    """在截图上标注操作步骤信息。

    策略：
    1. 先尝试用视觉 LLM 检测 UI 元素坐标并画边框（精确标注）
    2. 无论是否检测成功，都在底部添加操作步骤信息条（确保可读性）
    """
    # 收集 click/query/input/select 类型步骤
    action_keywords = {"click", "query", "select", "input", "login", "open_url"}
    action_steps = []
    for step in proc.steps:
        if step.action.lower() in action_keywords and step.target:
            action_steps.append(step)

    if not action_steps:
        return _normalize_image(frame_path)

    img = PILImage.open(frame_path).convert("RGB")
    w, h = img.size

    # 加载字体
    font_path = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"
    try:
        font_label = ImageFont.truetype(font_path, max(13, min(w, h) // 45))
        font_bar = ImageFont.truetype(font_path, max(12, min(w, h) // 50))
    except Exception:
        font_label = ImageFont.load_default()
        font_bar = font_label

    draw = ImageDraw.Draw(img, "RGBA")

    # === 策略1: 尝试 VLM bbox 检测（需要设置 ENABLE_BBOX_DETECTION=1）===
    boxes = []
    if os.getenv("ENABLE_BBOX_DETECTION", "").strip() == "1":
        targets = [s.target.strip() for s in action_steps[:3]]
        boxes = _detect_ui_targets(frame_path, targets)

    colors = [
        ((220, 40, 40), (220, 40, 40, 50)),     # 红
        ((30, 100, 220), (30, 100, 220, 50)),    # 蓝
        ((230, 140, 0), (230, 140, 0, 50)),      # 橙
    ]

    bbox_drawn = 0
    target_to_step = {s.target.strip(): s for s in action_steps}
    for idx, box in enumerate(boxes):
        bbox = box["bbox"]
        target_name = box.get("target", "")
        x1, y1, x2, y2 = bbox

        # 确保坐标在图片范围内
        x1, y1 = max(0, x1), max(0, y1)
        x2, y2 = min(w, x2), min(h, y2)
        if x2 - x1 < 5 or y2 - y1 < 5:
            continue

        color_border, color_fill = colors[idx % len(colors)]
        lw = max(2, min(w, h) // 250)

        # 半透明填充 + 边框
        draw.rectangle([x1, y1, x2, y2], fill=color_fill)
        for off in range(lw):
            draw.rectangle([x1 - off, y1 - off, x2 + off, y2 + off], outline=color_border)

        # 标签
        step = target_to_step.get(target_name)
        label = f"Step {step.no}" if step else target_name[:8]
        tb = draw.textbbox((0, 0), label, font=font_label)
        tw, th = tb[2] - tb[0] + 10, tb[3] - tb[1] + 6
        ly = y1 - th - 2
        if ly < 0:
            ly = y2 + 2
        draw.rectangle([x1, ly, x1 + tw, ly + th], fill=color_border)
        draw.text((x1 + 5, ly + 2), label, fill=(255, 255, 255), font=font_label)
        bbox_drawn += 1

    # === 策略2: 底部信息条（始终添加）===
    # 构建步骤文本
    step_texts = []
    action_cn = {"click": "点击", "query": "查询", "input": "输入", "select": "选择",
                 "login": "登录", "open_url": "打开"}
    for s in action_steps[:4]:
        act = action_cn.get(s.action.lower(), s.action)
        text = f"[{s.no}]{act}-{s.target}"
        if len(text) > 20:
            text = text[:19] + "..."
        step_texts.append(text)
    bar_text = "  ".join(step_texts)

    # 计算信息条高度
    tb = draw.textbbox((0, 0), bar_text, font=font_bar)
    bar_h = tb[3] - tb[1] + 14

    # 扩展画布高度（底部加信息条）
    new_img = PILImage.new("RGB", (w, h + bar_h), (255, 255, 255))
    new_img.paste(img, (0, 0))
    draw2 = ImageDraw.Draw(new_img, "RGBA")

    # 绘制信息条背景
    bar_bg = (40, 60, 90)
    draw2.rectangle([0, h, w, h + bar_h], fill=bar_bg)
    draw2.text((8, h + 5), bar_text, fill=(255, 255, 255), font=font_bar)

    buf = io.BytesIO()
    new_img.save(buf, format="PNG")
    buf.seek(0)
    found_msg = f", bbox标注{bbox_drawn}个" if bbox_drawn else ""
    logger.info(f"截图标注完成 [{proc.name}]: {len(action_steps)}个操作步骤{found_msg}")
    return buf


def generate_spec_doc(
    requirement_id: str,
    title: str,
    form_info: dict,
    extraction: ExtractionResult,
    screenshot_paths: Optional[List[Union[str, Tuple[str, float]]]] = None,
    template_path: Optional[str] = None,
    frame_descriptions: Optional[dict] = None,
) -> str:
    """生成 RPA 需求规格说明书 Word 文档"""
    if template_path is None:
        template_path = str(TEMPLATE_DIR / "rpa_spec_template.docx")

    if not os.path.exists(template_path):
        logger.warning(f"模板文件不存在: {template_path}，使用纯 python-docx 生成")
        return _generate_without_template(
            requirement_id, title, form_info, extraction, screenshot_paths,
            frame_descriptions=frame_descriptions,
        )

    logger.info(f"使用模板生成文档: {template_path}")
    doc = DocxTemplate(template_path)

    # 构建模板上下文
    context = _build_context(title, form_info, extraction)

    # 插入截图
    if screenshot_paths:
        images = []
        for i, item in enumerate(screenshot_paths):
            path = item[0] if isinstance(item, tuple) else item
            if os.path.exists(path):
                try:
                    img_buf = _normalize_image(path)
                    img = InlineImage(doc, img_buf, width=Mm(140))
                    images.append({"no": i + 1, "image": img, "desc": f"步骤 {i + 1} 截图"})
                except Exception as e:
                    logger.warning(f"插入截图失败 {path}: {e}")
        context["screenshots"] = images

    doc.render(context)
    output_path = str(OUTPUT_DIR / f"{requirement_id}_spec.docx")
    doc.save(output_path)
    logger.info(f"文档已生成: {output_path}")
    return output_path


def _match_frames_to_process(
    proc,
    timed_frames: List[Tuple[str, float]],
    all_procs: list,
) -> List[Tuple[str, float]]:
    """根据时间范围将帧匹配到对应的流程

    优先使用 LLM 输出的 ts_start_seconds/ts_end_seconds；
    若 LLM 未输出时间范围，则按流程顺序均分全部帧。
    """
    if not timed_frames:
        return []

    ts_start = getattr(proc, 'ts_start_seconds', None)
    ts_end = getattr(proc, 'ts_end_seconds', None)

    # 检查帧是否带有时间戳（timestamp >= 0 表示有）
    has_timestamps = any(ts >= 0 for _, ts in timed_frames)

    if ts_start is not None and ts_end is not None and has_timestamps:
        # 按时间范围精确匹配
        matched = [(p, t) for p, t in timed_frames if ts_start <= t <= ts_end]
        if matched:
            return matched
        # 如果精确匹配为空，找时间范围中点最近的帧
        mid = (ts_start + ts_end) / 2
        sorted_by_dist = sorted(timed_frames, key=lambda x: abs(x[1] - mid))
        # 返回最近的 1-2 帧
        return sorted_by_dist[:min(2, len(sorted_by_dist))]

    # 回退：按流程顺序均分
    proc_idx = all_procs.index(proc)
    total_procs = len(all_procs)
    per_proc = max(1, len(timed_frames) // total_procs)
    start_idx = proc_idx * per_proc
    end_idx = start_idx + per_proc if proc_idx < total_procs - 1 else len(timed_frames)
    return timed_frames[start_idx:end_idx]


def _filter_frames_by_relevance(
    proc,
    candidate_frames: List[Tuple[str, float]],
    frame_descriptions: dict,
    max_frames: int = 3,
) -> List[Tuple[str, float]]:
    """调用 LLM 从候选帧中筛选与流程内容相关的截图。

    Args:
        proc: MainProcess 对象（包含 name, steps）
        candidate_frames: 时间戳匹配后的候选帧 [(path, timestamp), ...]
        frame_descriptions: {timestamp_float: "视觉描述文本", ...}
        max_frames: 最多保留帧数

    Returns:
        筛选后的帧列表（仅保留与流程内容相关的帧）
    """
    if not candidate_frames or not frame_descriptions:
        return candidate_frames[:max_frames]

    # 构建候选帧描述列表（截断描述以减少 token 消耗）
    candidates_text = []
    for i, (path, ts) in enumerate(candidate_frames):
        desc = frame_descriptions.get(ts, "")
        if not desc:
            # 尝试模糊匹配（浮点精度）
            for k, v in frame_descriptions.items():
                if abs(k - ts) < 1.0:
                    desc = v
                    break
        # 截断过长描述
        if desc and len(desc) > 40:
            desc = desc[:40]
        ts_label = f"{int(ts//60):02d}:{int(ts%60):02d}"
        candidates_text.append(f"[{i}] {ts_label} {desc or '无描述'}")

    # 构建流程步骤描述（简洁格式）
    steps_text = " ".join(
        f"{s.no}.{s.target}"
        for s in proc.steps
    )

    prompt = f"""判断哪些截图与流程相关，仅输出JSON。

流程: {proc.name}
步骤: {steps_text}

截图:
{chr(10).join(candidates_text)}

输出: {{"relevant_indices": [相关截图编号], "reason": "理由"}}
最多保留{max_frames}张最相关的。"""

    try:
        api_key = os.getenv("LLM_API_KEY", "")
        base_url = os.getenv("LLM_BASE_URL", "")
        model = os.getenv("LLM_MODEL", "")

        if not all([api_key, base_url, model]):
            logger.warning("LLM 配置缺失，跳过相关性过滤")
            return candidate_frames[:max_frames]

        client = OpenAI(
            api_key=api_key,
            base_url=base_url,
            timeout=httpx.Timeout(60.0, connect=10.0),
        )
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=2048,
        )
        content = (response.choices[0].message.content or "").strip()
        # 移除 <think>...</think>
        content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
        # 提取 JSON
        json_match = re.search(r"\{.*\}", content, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())
            indices = result.get("relevant_indices", [])
            reason = result.get("reason", "")
            logger.info(f"帧相关性过滤 [{proc.name}]: {len(candidate_frames)}→{len(indices)} 帧, {reason}")
            filtered = [candidate_frames[i] for i in indices if 0 <= i < len(candidate_frames)]
            return filtered[:max_frames] if filtered else []
        else:
            logger.warning(f"帧相关性过滤解析失败: {content[:200]}")
            return candidate_frames[:max_frames]

    except Exception as e:
        logger.warning(f"帧相关性过滤调用失败（降级保留全部）: {e}")
        return candidate_frames[:max_frames]


def _generate_flowchart(process_names: List[str]) -> Optional[io.BytesIO]:
    """使用 matplotlib 生成垂直流程图，返回 PNG 图片 BytesIO"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        from matplotlib.font_manager import FontProperties

        font_path = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"
        if not os.path.exists(font_path):
            logger.warning(f"流程图字体不存在: {font_path}，跳过流程图")
            return None
        fp = FontProperties(fname=font_path, size=11)
        fp_small = FontProperties(fname=font_path, size=9)

        n = len(process_names)
        if n == 0:
            return None

        # 布局参数
        box_w, box_h = 3.6, 0.55
        gap = 0.35
        oval_h = 0.45
        total_h = oval_h + gap + n * (box_h + gap) + oval_h
        fig_w = box_w + 1.6
        fig_h = max(total_h + 0.6, 3.0)

        fig, ax = plt.subplots(figsize=(fig_w, fig_h))
        ax.set_xlim(0, fig_w)
        ax.set_ylim(0, fig_h)
        ax.axis("off")

        cx = fig_w / 2
        y = fig_h - 0.4  # top

        # --- 开始 椭圆 ---
        start_y = y - oval_h / 2
        oval_start = mpatches.FancyBboxPatch(
            (cx - box_w / 4, start_y - oval_h / 2), box_w / 2, oval_h,
            boxstyle="round,pad=0.15", facecolor="#4CAF50", edgecolor="#388E3C", linewidth=1.5,
        )
        ax.add_patch(oval_start)
        ax.text(cx, start_y, "开始", ha="center", va="center",
                fontproperties=fp, color="white", fontweight="bold")
        y = start_y - oval_h / 2

        prev_bottom = y

        # --- 流程节点 ---
        for i, name in enumerate(process_names):
            y -= gap
            rect_top = y
            rect_bottom = y - box_h
            rect = mpatches.FancyBboxPatch(
                (cx - box_w / 2, rect_bottom), box_w, box_h,
                boxstyle="round,pad=0.08", facecolor="#E3F2FD", edgecolor="#1565C0", linewidth=1.2,
            )
            ax.add_patch(rect)
            # 截断过长名称
            label = name if len(name) <= 16 else name[:15] + "…"
            ax.text(cx, (rect_top + rect_bottom) / 2, f"{i + 1}. {label}",
                    ha="center", va="center", fontproperties=fp_small, color="#0D47A1")
            # 箭头
            ax.annotate("", xy=(cx, rect_top), xytext=(cx, prev_bottom),
                        arrowprops=dict(arrowstyle="-|>", color="#555555", lw=1.3))
            prev_bottom = rect_bottom
            y = rect_bottom

        # --- 结束 椭圆 ---
        y -= gap
        end_y = y - oval_h / 2
        oval_end = mpatches.FancyBboxPatch(
            (cx - box_w / 4, end_y - oval_h / 2), box_w / 2, oval_h,
            boxstyle="round,pad=0.15", facecolor="#F44336", edgecolor="#C62828", linewidth=1.5,
        )
        ax.add_patch(oval_end)
        ax.text(cx, end_y, "结束", ha="center", va="center",
                fontproperties=fp, color="white", fontweight="bold")
        ax.annotate("", xy=(cx, end_y + oval_h / 2), xytext=(cx, prev_bottom),
                    arrowprops=dict(arrowstyle="-|>", color="#555555", lw=1.3))

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        logger.info(f"流程图已生成（{n} 个流程节点）")
        return buf
    except Exception as e:
        logger.warning(f"生成流程图失败: {e}")
        return None


def _generate_without_template(
    requirement_id: str,
    title: str,
    form_info: dict,
    extraction: ExtractionResult,
    screenshot_paths: Optional[List[Union[str, Tuple[str, float]]]] = None,
    frame_descriptions: Optional[dict] = None,
) -> str:
    """不使用模板，直接用 python-docx 生成文档（对齐 RPA 需求采集模板 13 章结构）"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ========== 设置默认字体 ==========
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)

    # ========== 封面 ==========
    doc.add_paragraph("")
    doc.add_paragraph("")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f"{title or 'RPA 需求规格说明书'}")
    run.font.size = Pt(22)
    run.font.bold = True

    doc.add_paragraph("")
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("RPA 需求规格说明书")
    run.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 文档信息表
    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("文档编号", requirement_id[:8].upper()),
        ("版本号", "V1.0"),
        ("编写日期", datetime.now().strftime("%Y-%m-%d")),
        ("需求部门", form_info.get("req_dept", "")),
        ("需求提出人", form_info.get("req_owner", "")),
        ("优先级", form_info.get("priority", "中")),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v

    doc.add_page_break()

    # ========== 1. 需求概述 ==========
    doc.add_heading("1. 需求概述", level=1)

    doc.add_heading("1.1 基础信息", level=2)
    base_info_table = doc.add_table(rows=1, cols=2)
    base_info_table.style = 'Table Grid'
    hdr = base_info_table.rows[0].cells
    hdr[0].text = "字段"
    hdr[1].text = "内容"
    base_fields = [
        ("需求名称", title or "—"),
        ("需求部门", form_info.get("req_dept", "—")),
        ("需求联系人", form_info.get("req_owner", "—")),
        ("联系方式", form_info.get("contact_info", "—")),
        ("需求类型", form_info.get("req_type", "—")),
        ("优先级", form_info.get("priority", "中")),
        ("执行频率", form_info.get("exec_frequency", "—")),
        ("当前执行角色", form_info.get("current_role", "—")),
    ]
    for k, v in base_fields:
        row = base_info_table.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_heading("1.2 自动化目标", level=2)
    overview = extraction.business_overview
    doc.add_paragraph(overview.auto_goal if overview.auto_goal else "—")

    doc.add_heading("1.3 业务范围", level=2)
    doc.add_paragraph(overview.scope if overview.scope else "—")

    # ========== 2. 业务背景与痛点 ==========
    doc.add_heading("2. 业务背景与痛点", level=1)

    doc.add_heading("2.1 需求背景", level=2)
    doc.add_paragraph(form_info.get("req_background", "") or "—")

    doc.add_heading("2.2 当前痛点", level=2)
    doc.add_paragraph(form_info.get("current_pain", "") or "—")

    doc.add_heading("2.3 业务量", level=2)
    vol_fields = [
        ("单次耗时", form_info.get("single_duration", "—")),
        ("业务量", form_info.get("business_volume", "—")),
        ("日常执行时段", form_info.get("execution_time", "—")),
        ("建议 RPA 执行时间", form_info.get("rpa_schedule_time", "—")),
    ]
    for label, val in vol_fields:
        doc.add_paragraph(f"{label}：{val}")

    # ========== 3. 当前人工流程说明 ==========
    doc.add_heading("3. 当前人工流程说明", level=1)
    manual_desc = extraction.manual_flow_description
    if manual_desc:
        doc.add_paragraph(manual_desc)
    else:
        doc.add_paragraph("（由 AI 从视频内容分析生成，如需补充请编辑。）")

    # ========== 4. 涉及系统与对象 ==========
    doc.add_heading("4. 涉及系统与对象", level=1)
    if extraction.system_env:
        sys_table = doc.add_table(rows=1, cols=3)
        sys_table.style = 'Table Grid'
        hdr = sys_table.rows[0].cells
        hdr[0].text = "系统名称"
        hdr[1].text = "认证方式"
        hdr[2].text = "浏览器要求"
        for env in extraction.system_env:
            row = sys_table.add_row().cells
            row[0].text = env.name
            row[1].text = env.auth or "—"
            row[2].text = env.browser or "—"
    else:
        doc.add_paragraph("—")

    if form_info.get("involved_systems"):
        doc.add_paragraph("")
        doc.add_paragraph(f"涉及系统补充说明：{form_info['involved_systems']}")

    if form_info.get("target_url"):
        doc.add_paragraph(f"目标系统 URL：{form_info['target_url']}")

    doc.add_paragraph(f"是否需要登录：{'是' if form_info.get('login_required') else '否'}")

    # ========== 5. 流程步骤说明 ==========
    doc.add_heading("5. 流程步骤说明", level=1)

    # 插入流程概览图
    if extraction.main_process:
        flowchart_buf = _generate_flowchart([p.name for p in extraction.main_process])
        if flowchart_buf:
            p_fc = doc.add_paragraph()
            p_fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_fc = p_fc.add_run()
            run_fc.add_picture(flowchart_buf, width=Inches(4.5))
            cap_fc = doc.add_paragraph()
            cap_fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_fc.add_run("图: 整体流程概览")
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph("")  # 空行分隔

    # 解析截图：支持 (path, timestamp) 元组或纯 path 字符串
    timed_frames: List[Tuple[str, float]] = []
    if screenshot_paths:
        for item in screenshot_paths:
            if isinstance(item, tuple):
                timed_frames.append(item)
            else:
                timed_frames.append((item, -1.0))

    # ===== 预计算：并发执行所有流程的帧匹配 + LLM 过滤 + 标注 =====
    import time as _time
    from concurrent.futures import ThreadPoolExecutor, as_completed

    t_filter_start = _time.time()
    # 1) 先做帧匹配（CPU快速操作）
    proc_matched_frames = {}
    for proc in extraction.main_process:
        matched = _match_frames_to_process(proc, timed_frames, extraction.main_process)
        proc_matched_frames[proc.name] = matched

    # 2) 并发 LLM 过滤（I/O 密集型）
    if frame_descriptions:
        filter_concurrency = int(os.getenv("FILTER_CONCURRENCY", "3"))

        def _filter_one(proc_name, proc, matched):
            return proc_name, _filter_frames_by_relevance(
                proc, matched, frame_descriptions, max_frames=3,
            )

        with ThreadPoolExecutor(max_workers=filter_concurrency, thread_name_prefix="filter") as pool:
            futures = []
            for proc in extraction.main_process:
                matched = proc_matched_frames[proc.name]
                if matched:
                    futures.append(pool.submit(_filter_one, proc.name, proc, matched))
            for future in as_completed(futures):
                proc_name, filtered = future.result()
                proc_matched_frames[proc_name] = filtered

    # 3) 并发标注截图（I/O + 可能的 VLM 调用）
    # 收集所有需要标注的 (frame_path, proc) 组合
    annotate_tasks = []
    for proc in extraction.main_process:
        matched = proc_matched_frames.get(proc.name, [])
        for frame_path, frame_ts in matched[:3]:
            if os.path.exists(frame_path):
                annotate_tasks.append((frame_path, proc, frame_ts))

    annotated_images = {}
    if annotate_tasks:
        annotate_concurrency = int(os.getenv("ANNOTATE_CONCURRENCY", "3"))

        def _annotate_one(frame_path, proc):
            return (frame_path, proc.name), _annotate_frame(frame_path, proc, frame_descriptions)

        with ThreadPoolExecutor(max_workers=annotate_concurrency, thread_name_prefix="annotate") as pool:
            futures = [pool.submit(_annotate_one, fp, p) for fp, p, _ in annotate_tasks]
            for future in as_completed(futures):
                key, img_buf = future.result()
                annotated_images[key] = img_buf

    t_filter_elapsed = _time.time() - t_filter_start
    logger.info(f"帧过滤+标注并发耗时: {t_filter_elapsed:.1f}s (过滤{len(extraction.main_process)}流程, 标注{len(annotate_tasks)}张)")

    img_counter = 0
    for proc in extraction.main_process:
        proc_idx = extraction.main_process.index(proc) + 1
        doc.add_heading(f"5.{proc_idx} {proc.name}", level=2)

        step_table = doc.add_table(rows=1, cols=4)
        step_table.style = 'Table Grid'
        hdr = step_table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "操作类型"
        hdr[2].text = "操作目标"
        hdr[3].text = "参数/值"

        for step in proc.steps:
            row = step_table.add_row().cells
            row[0].text = str(step.no)
            row[1].text = step.action
            row[2].text = step.target
            row[3].text = step.value or step.result_file or "—"

        # 使用预计算的匹配+过滤结果
        matched_frames = proc_matched_frames.get(proc.name, [])
        max_per_proc = 3
        for frame_path, frame_ts in matched_frames[:max_per_proc]:
            if os.path.exists(frame_path):
                img_counter += 1
                doc.add_paragraph("")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    # 使用预计算的标注图
                    key = (frame_path, proc.name)
                    img_buf = annotated_images.get(key)
                    if img_buf is None:
                        img_buf = _annotate_frame(frame_path, proc, frame_descriptions)
                    run = p.add_run()
                    run.add_picture(img_buf, width=Inches(5.0))
                except Exception as e:
                    logger.warning(f"插入截图失败 {frame_path}: {e}")
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ts_label = f" ({int(frame_ts//60):02d}:{int(frame_ts%60):02d})" if frame_ts >= 0 else ""
                cap_run = cap.add_run(f"图 {img_counter}: {proc.name}{ts_label}")
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ========== 6. 业务规则说明 ==========
    doc.add_heading("6. 业务规则说明", level=1)
    if extraction.rules:
        for i, rule in enumerate(extraction.rules, 1):
            doc.add_paragraph(f"{i}. {rule}")
    else:
        doc.add_paragraph("暂无特殊业务规则。")

    # ========== 7. 异常与边界场景 ==========
    doc.add_heading("7. 异常与边界场景", level=1)
    if extraction.exceptions:
        exc_table = doc.add_table(rows=1, cols=2)
        exc_table.style = 'Table Grid'
        hdr = exc_table.rows[0].cells
        hdr[0].text = "异常代码"
        hdr[1].text = "处理方式"
        for exc in extraction.exceptions:
            row = exc_table.add_row().cells
            row[0].text = exc.code
            row[1].text = exc.handler
    else:
        doc.add_paragraph("暂无异常处理策略。")

    if form_info.get("exception_policy"):
        doc.add_paragraph("")
        doc.add_paragraph(f"异常通知方式：{'、'.join(form_info['exception_policy'])}")

    # ========== 8. 输入输出定义 ==========
    doc.add_heading("8. 输入输出定义", level=1)

    doc.add_heading("8.1 输入", level=2)
    if extraction.io_spec.input:
        for item in extraction.io_spec.input:
            doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("—")
    if form_info.get("input_source"):
        doc.add_paragraph(f"输入来源：{form_info['input_source']}")

    doc.add_heading("8.2 输出", level=2)
    if extraction.io_spec.output:
        for item in extraction.io_spec.output:
            doc.add_paragraph(f"• {item}")
    else:
        doc.add_paragraph("—")
    if form_info.get("output_sink"):
        doc.add_paragraph(f"输出去向：{form_info['output_sink']}")

    # ========== 9. 前置条件与依赖 ==========
    doc.add_heading("9. 前置条件与依赖", level=1)
    prereqs_written = False
    if extraction.prerequisites:
        for i, p in enumerate(extraction.prerequisites, 1):
            doc.add_paragraph(f"{i}. {p}")
        prereqs_written = True
    form_prereqs = [
        ("数据源前提", form_info.get("data_prerequisite")),
        ("目标系统前提", form_info.get("system_prerequisite")),
        ("其他依赖", form_info.get("other_dependency")),
    ]
    for label, val in form_prereqs:
        if val:
            doc.add_paragraph(f"{label}：{val}")
            prereqs_written = True
    if not prereqs_written:
        doc.add_paragraph("暂无特殊前置条件。")

    # ========== 10. 权限与安全要求 ==========
    doc.add_heading("10. 权限与安全要求", level=1)
    security_written = False

    # 账号信息
    acct_fields = [
        ("账号类型", form_info.get("account_type")),
        ("是否多人共用", "是" if form_info.get("multi_user") else None),
        ("操作权限限制", form_info.get("permission_limit")),
        ("是否涉及敏感数据", "是" if form_info.get("sensitive_data") else None),
        ("合规要求", form_info.get("compliance_req")),
    ]
    for label, val in acct_fields:
        if val:
            doc.add_paragraph(f"{label}：{val}")
            security_written = True

    if extraction.security_requirements:
        for i, s in enumerate(extraction.security_requirements, 1):
            doc.add_paragraph(f"{i}. {s}")
        security_written = True

    if not security_written:
        doc.add_paragraph("—")

    # ========== 11. 自动化预期收益 ==========
    doc.add_heading("11. 自动化预期收益", level=1)
    benefit_table = doc.add_table(rows=1, cols=2)
    benefit_table.style = 'Table Grid'
    hdr = benefit_table.rows[0].cells
    hdr[0].text = "指标"
    hdr[1].text = "说明"
    benefit_fields = [
        ("当前投入人力", form_info.get("current_headcount", "—")),
        ("当前工时", form_info.get("current_hours", "—")),
        ("预期收益", form_info.get("expected_benefit", "—")),
        ("预期节省工时", form_info.get("expected_saving", "—")),
        ("质量改进", form_info.get("quality_improvement", "—")),
    ]
    for k, v in benefit_fields:
        row = benefit_table.add_row().cells
        row[0].text = k
        row[1].text = v

    # ========== 12. 可行性初步判断 ==========
    doc.add_heading("12. 可行性初步判断", level=1)
    if extraction.feasibility_notes:
        for i, note in enumerate(extraction.feasibility_notes, 1):
            doc.add_paragraph(f"{i}. {note}")
    else:
        doc.add_paragraph("—")

    # ========== 13. 待确认问题清单 ==========
    doc.add_heading("13. 待确认问题清单", level=1)
    if extraction.pending_questions:
        q_table = doc.add_table(rows=1, cols=3)
        q_table.style = 'Table Grid'
        hdr = q_table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "待确认问题"
        hdr[2].text = "状态"
        for i, q in enumerate(extraction.pending_questions, 1):
            row = q_table.add_row().cells
            row[0].text = str(i)
            row[1].text = q
            row[2].text = "待确认"
    else:
        doc.add_paragraph("暂无待确认问题。")

    # ========== 附录：运行环境 ==========
    doc.add_heading("附录：运行环境", level=1)
    env_fields = [
        ("电脑配置", form_info.get("pc_config", "—")),
        ("浏览器", form_info.get("browser", "—")),
        ("网络环境", form_info.get("network_env", "—")),
    ]
    for label, val in env_fields:
        doc.add_paragraph(f"{label}：{val or '—'}")

    # 保存
    output_path = str(OUTPUT_DIR / f"{requirement_id}_spec.docx")
    doc.save(output_path)
    logger.info(f"文档已生成（无模板模式）: {output_path}")
    return output_path


def _build_context(title: str, form_info: dict, extraction: ExtractionResult) -> dict:
    """构建 docxtpl 模板上下文"""
    now = datetime.now()
    return {
        "doc_title": title or "RPA 需求规格说明书",
        "doc_id": form_info.get("requirement_id", "")[:8].upper(),
        "doc_version": "V1.0",
        "doc_date": now.strftime("%Y-%m-%d"),
        "req_dept": form_info.get("req_dept", ""),
        "req_owner": form_info.get("req_owner", ""),
        "req_type": form_info.get("req_type", ""),
        "target_url": form_info.get("target_url", ""),
        "login_required": "是" if form_info.get("login_required") else "否",
        "exec_frequency": form_info.get("exec_frequency", ""),
        "auto_goal": extraction.business_overview.auto_goal,
        "scope": extraction.business_overview.scope,
        "req_background": form_info.get("req_background", ""),
        "current_pain": form_info.get("current_pain", ""),
        "current_role": form_info.get("current_role", ""),
        "priority": form_info.get("priority", "中"),
        "manual_flow_description": extraction.manual_flow_description,
        "main_process": [
            {
                "name": p.name,
                "steps": [
                    {
                        "no": s.no,
                        "action": s.action,
                        "target": s.target,
                        "value": s.value or s.result_file or "",
                    }
                    for s in p.steps
                ],
            }
            for p in extraction.main_process
        ],
        "io_input": extraction.io_spec.input,
        "io_output": extraction.io_spec.output,
        "rules": extraction.rules,
        "system_env": [
            {"name": e.name, "auth": e.auth or "", "browser": e.browser or ""}
            for e in extraction.system_env
        ],
        "exceptions": [
            {"code": e.code, "handler": e.handler}
            for e in extraction.exceptions
        ],
        "prerequisites": extraction.prerequisites,
        "security_requirements": extraction.security_requirements,
        "feasibility_notes": extraction.feasibility_notes,
        "pending_questions": extraction.pending_questions,
    }
